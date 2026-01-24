#!/usr/bin/env python3
"""
Extract ALL G2 questions from Word documents - IMPROVED VERSION
Handles multiple document formats:
- Questions with/without numbering
- Answers marked with underline formatting
- Questions with 4-5 options
- Various punctuation patterns
"""

import json
import os
import re
from docx import Document
from difflib import SequenceMatcher

def normalize_text(text):
    """Normalize text for comparison"""
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text.strip())
    text = re.sub(r'[^\w\s]', '', text)
    return text.lower()

def calculate_similarity(str1, str2):
    """Calculate similarity ratio between two strings"""
    norm1 = normalize_text(str1)
    norm2 = normalize_text(str2)
    return SequenceMatcher(None, norm1, norm2).ratio()

def extract_questions_from_docx(file_path):
    """
    Extract questions from Word document - handles multiple formats
    """
    print(f"  Reading: {os.path.basename(file_path)}")

    try:
        doc = Document(file_path)
        questions = []
        i = 0

        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            text = para.text.strip()

            # Check if this is a question
            # Must be substantial length and end with ? or :
            is_question = (text and
                          len(text) > 20 and
                          (text.endswith('?') or text.endswith(':')))

            # Also accept numbered questions like "1. Question text?"
            if not is_question and text:
                question_match = re.match(r'^\d+[\.\)]\s*(.+[?:])$', text)
                if question_match:
                    text = question_match.group(1).strip()
                    is_question = True

            if is_question:
                # Collect options from next paragraphs
                options = []
                correct_index = None
                j = i + 1

                while j < len(doc.paragraphs) and j < i + 12:
                    opt_para = doc.paragraphs[j]
                    opt_text = opt_para.text.strip()

                    # Skip empty paragraphs
                    if not opt_text:
                        j += 1
                        continue

                    # Stop if we hit another question
                    if len(opt_text) > 20 and (opt_text.endswith('?') or opt_text.endswith(':')):
                        break

                    # Stop if we hit a Note paragraph (explanatory text)
                    if opt_text.startswith('Note'):
                        break

                    # Check if this is an option (short text, < 200 chars)
                    if len(opt_text) < 200:
                        # Clean option text - remove leading a), b), etc. if present
                        cleaned_opt = re.sub(r'^[a-e][\.\)]\s*', '', opt_text, flags=re.IGNORECASE)

                        # Check if this option is underlined (correct answer)
                        is_underlined = any(run.underline for run in opt_para.runs if run.text.strip())
                        if is_underlined and correct_index is None:
                            correct_index = len(options)

                        options.append(cleaned_opt)

                        # If we have 4-5 options and one is underlined, we're done
                        if len(options) >= 4 and correct_index is not None:
                            break
                        # If we have 5+ options and no underline detected, take first 4
                        if len(options) >= 5 and correct_index is None:
                            options = options[:4]
                            break
                    else:
                        # Long text likely means we've moved past options
                        break

                    j += 1

                # Only accept questions with 4 options
                if len(options) == 4:
                    questions.append({
                        'question': text,
                        'options': options,
                        'correct': correct_index if correct_index is not None else 0,
                        'source': os.path.basename(file_path),
                        'needs_verification': correct_index is None
                    })
                    i = j
                else:
                    i += 1
            else:
                i += 1

        verified = sum(1 for q in questions if not q.get('needs_verification'))
        print(f"    Found {len(questions)} questions ({verified} with verified answers)")
        return questions

    except Exception as e:
        print(f"    ERROR: {e}")
        return []

def is_duplicate(new_q, existing_questions, similarity_threshold=0.90):
    """
    Check if question is a duplicate
    - Exact matches (>90% similar): duplicate
    - Variations (70-90% similar): keep both
    - Different (<70% similar): not duplicate
    """
    new_text = normalize_text(new_q['question'])

    for existing_q in existing_questions:
        existing_text = normalize_text(existing_q['question'])
        similarity = calculate_similarity(new_text, existing_text)

        if similarity >= similarity_threshold:
            return True, existing_q

    return False, None

def categorize_question(question_text):
    """Categorize question by CSA unit using keyword matching"""
    q_lower = question_text.lower()

    unit_keywords = {
        10: ['pipe', 'piping', 'weld', 'tube', 'tubing', 'purge', 'nps', 'schedule', 'elbow', 'tee', 'fitting', 'underground'],
        11: ['regulator', 'meter', 'relief', 'propane cylinder', 'tank', 'container', 'diaphragm', 'pressure regulator', 'overpressure'],
        12: ['electrical', 'fuse', 'wire', 'amperage', 'voltage', 'resistance', 'ohm', 'circuit', 'transformer', 'awg', 'thermostat', 'watt'],
        13: ['gas valve', 'thermocouple', 'thermopile', 'hot surface', 'flame rod', 'pilot', 'aquastat', 'limit', 'control', 'ignition'],
        14: ['building', 'combustion air', 'ventilation', 'depressurization', 'stack effect', 'distribution', 'air supply', 'infiltration'],
        15: ['range', 'dryer', 'clothes', 'barbecue', 'lamp', 'oven', 'moisture exhaust', 'flexible connector', 'appliance'],
        16: ['refrigerator', 'absorption', 'ammonia', 'generator', 'evaporator', 'condenser'],
        17: ['conversion', 'converting', 'flue gas analysis', 'convert', 'orifice'],
        18: ['water heater', 'combination system', 'circulator', 'mixing valve', 'dip tube', 'anode', 'storage tank', 'relief valve', 'indirect'],
        19: ['furnace', 'forced air', 'blower', 'heat exchanger', 'high limit', 'filter', 'duct', 'belt', 'pulley', 'plenum', 'warm air'],
        20: ['boiler', 'hydronic', 'expansion tank', 'lwco', 'backflow', 'pool heater', 'steam', 'fire tube', 'water tube', 'heating water'],
        21: ['space heater', 'fireplace', 'radiant', 'infra-red', 'infrared', 'decorative', 'room heater', 'unvented'],
        22: ['vent', 'venting', 'chimney', 'draft', 'flue', 'b vent', 'spillage', 'connector', 'barometric', 'drafthood'],
        23: ['humidifier', 'electronic air', 'air filtration', 'humidistat', 'ionizing', 'air cleaner', 'filter'],
        24: ['heat loss', 'heat gain', 'btu', 'cfm', 'air handling', 'convection', 'sensible', 'latent', 'afue', 'efficiency']
    }

    unit_scores = {}
    for unit, keywords in unit_keywords.items():
        score = sum(1 for kw in keywords if kw in q_lower)
        if score > 0:
            unit_scores[unit] = score

    if unit_scores:
        return max(unit_scores.items(), key=lambda x: x[1])[0]

    return None

def main():
    print("=" * 70)
    print("EXTRACTING ALL G2 QUESTIONS - IMPROVED VERSION")
    print("=" * 70)

    exam_prep_dir = r"C:\LocalProjects\WorkFlow\TSSA G2 Exam Prep"

    # Load existing questions
    print("\n[1] Loading existing 155 questions from simulator...")
    existing_path = r"C:\LocalProjects\WorkFlow\hvac-student-resources\g2_question_bank_final.js"
    with open(existing_path, 'r', encoding='utf-8') as f:
        js_content = f.read()

    existing_questions = []
    pattern = r'question: "([^"]+)".*?options: (\[[^\]]+\]).*?correct: (\d+)'
    for match in re.finditer(pattern, js_content, re.DOTALL):
        q_text = match.group(1)
        options_str = match.group(2)
        correct = int(match.group(3))
        options = eval(options_str)

        existing_questions.append({
            'question': q_text,
            'options': options,
            'correct': correct,
            'source': 'existing'
        })

    print(f"  Loaded {len(existing_questions)} existing questions")

    # Find all Word documents
    print(f"\n[2] Scanning for Word documents...")
    docx_files = []
    for file in os.listdir(exam_prep_dir):
        if file.endswith('.docx') and not file.startswith('~'):
            docx_files.append(os.path.join(exam_prep_dir, file))

    print(f"  Found {len(docx_files)} Word documents")

    # Extract questions
    print(f"\n[3] Extracting questions from Word documents...")
    all_extracted = []
    for doc_path in sorted(docx_files):
        questions = extract_questions_from_docx(doc_path)
        all_extracted.extend(questions)

    print(f"\n  Total extracted: {len(all_extracted)} questions")

    # Count verified vs unverified
    verified = sum(1 for q in all_extracted if not q.get('needs_verification'))
    unverified = len(all_extracted) - verified
    print(f"  Verified answers: {verified}")
    print(f"  Need verification: {unverified}")

    # Remove duplicates
    print(f"\n[4] Removing duplicates (keeping variations)...")
    unique_questions = []
    duplicates = 0
    variations = 0

    for new_q in all_extracted:
        is_dup, matched = is_duplicate(new_q, existing_questions + unique_questions, similarity_threshold=0.90)

        if is_dup:
            duplicates += 1
        else:
            is_variation = False
            for existing_q in existing_questions + unique_questions:
                similarity = calculate_similarity(new_q['question'], existing_q['question'])
                if 0.70 <= similarity < 0.90:
                    is_variation = True
                    variations += 1
                    break

            unique_questions.append(new_q)

    print(f"  Duplicates removed: {duplicates}")
    print(f"  Variations kept: {variations}")
    print(f"  New unique questions: {len(unique_questions)}")

    # Categorize new questions
    print(f"\n[5] Categorizing new questions by CSA units...")
    questions_by_unit = {}
    uncategorized = []

    for q in unique_questions:
        unit = categorize_question(q['question'])
        if unit:
            if unit not in questions_by_unit:
                questions_by_unit[unit] = []
            questions_by_unit[unit].append(q)
        else:
            uncategorized.append(q)

    print(f"  Categorized: {sum(len(v) for v in questions_by_unit.values())}")
    print(f"  Uncategorized: {len(uncategorized)}")

    for unit in sorted(questions_by_unit.keys()):
        print(f"    Unit {unit}: +{len(questions_by_unit[unit])} new questions")

    # Combine with existing questions
    print(f"\n[6] Merging with existing question bank...")

    # Load existing with units
    existing_with_units = {}
    pattern = r'\{\s*id: (\d+),\s*unit: (\d+),\s*question: "([^"]+)".*?\}'
    for match in re.finditer(pattern, js_content, re.DOTALL):
        qid = int(match.group(1))
        unit = int(match.group(2))
        q_text = match.group(3)

        if unit not in existing_with_units:
            existing_with_units[unit] = []

        for eq in existing_questions:
            if eq['question'] == q_text:
                eq['unit'] = unit
                eq['id'] = qid
                existing_with_units[unit].append(eq)
                break

    # Combine
    combined_by_unit = {}
    for unit in set(list(existing_with_units.keys()) + list(questions_by_unit.keys())):
        combined_by_unit[unit] = existing_with_units.get(unit, []) + questions_by_unit.get(unit, [])

    total_questions = sum(len(v) for v in combined_by_unit.values())

    print(f"  Total questions in combined bank: {total_questions}")
    print(f"  Previous: 155")
    print(f"  New additions: {total_questions - 155}")

    # Generate JavaScript
    print(f"\n[7] Generating JavaScript question bank...")
    js_lines = []
    js_lines.append("// TSSA G2 Comprehensive Question Bank")
    js_lines.append(f"// Total: {total_questions} questions from all sources")
    js_lines.append("// Sources: Official exam prep + 6 practice exams + 8 review documents")
    js_lines.append("// Organized by CSA Units 10-24")
    js_lines.append("")
    js_lines.append("const rawQuestionBank = [")

    question_id = 1
    for unit in sorted(combined_by_unit.keys()):
        js_lines.append(f"\n    // ========================================")
        js_lines.append(f"    // CSA UNIT {unit} - {len(combined_by_unit[unit])} Questions")
        js_lines.append(f"    // ========================================")

        for q in combined_by_unit[unit]:
            js_lines.append("    {")
            js_lines.append(f"        id: {question_id},")
            js_lines.append(f"        unit: {unit},")
            js_lines.append(f"        question: {json.dumps(q['question'], ensure_ascii=False)},")
            js_lines.append(f"        options: {json.dumps(q['options'], ensure_ascii=False)},")

            if q.get('correct') is not None:
                js_lines.append(f"        correct: {q['correct']},")
            else:
                js_lines.append(f"        correct: 0,  // NEEDS VERIFICATION")

            if q.get('needs_verification'):
                js_lines.append(f"        // WARNING: Answer needs verification")

            if q.get('source'):
                js_lines.append(f"        source: {json.dumps(q['source'], ensure_ascii=False)}")

            js_lines.append("    },")
            question_id += 1

    # Remove trailing comma
    if js_lines[-1].endswith(","):
        js_lines[-1] = js_lines[-1][:-1]

    js_lines.append("];")
    js_lines.append("")
    js_lines.append("// Question bank statistics")
    js_lines.append("const QUESTION_BANK_STATS = {")
    js_lines.append(f"    totalQuestions: {total_questions},")
    js_lines.append("    questionsByUnit: {")
    for unit in sorted(combined_by_unit.keys()):
        js_lines.append(f"        {unit}: {len(combined_by_unit[unit])},")
    js_lines.append("    }")
    js_lines.append("};")

    # Write output
    output_path = r"C:\LocalProjects\WorkFlow\hvac-student-resources\g2_question_bank_expanded.js"
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(js_lines))

    print(f"\n[SUCCESS] Generated: {output_path}")

    # Summary report
    print(f"\n" + "=" * 70)
    print("EXTRACTION SUMMARY")
    print("=" * 70)
    print(f"Documents processed: {len(docx_files)}")
    print(f"Questions extracted: {len(all_extracted)}")
    print(f"  - With verified answers: {verified}")
    print(f"  - Need verification: {unverified}")
    print(f"Duplicates removed: {duplicates}")
    print(f"Variations kept: {variations}")
    print(f"New unique additions: {len(unique_questions)}")
    print(f"Previous question count: 155")
    print(f"NEW TOTAL: {total_questions} questions")
    print(f"Increase: +{total_questions - 155} questions ({((total_questions - 155) / 155 * 100):.1f}%)")
    print("=" * 70)

if __name__ == "__main__":
    main()
